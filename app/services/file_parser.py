from typing import List, BinaryIO, Union
import logging
from pathlib import Path
import pandas as pd
from docx import Document
from app.models.document import (
    DocumentChunk,
    ParsedDocument,
    FileType,
    ExcelMetadata,
    WordMetadata,
    FileValidationError,
)

logger = logging.getLogger(__name__)

# File size limits (in bytes)
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
MAX_EXCEL_SIZE = 20 * 1024 * 1024  # 20 MB for Excel
MAX_WORD_SIZE = 10 * 1024 * 1024  # 10 MB for Word

# Supported file extensions
EXCEL_EXTENSIONS = {'.xlsx', '.xls', '.xlsm'}
WORD_EXTENSIONS = {'.docx', '.doc'}


def validate_file_type(filename: str, allowed_extensions: set) -> None:
    """
    Validate that the file has an allowed extension.

    Args:
        filename: Name of the file
        allowed_extensions: Set of allowed file extensions

    Raises:
        FileValidationError: If file type is not supported
    """
    file_ext = Path(filename).suffix.lower()
    if file_ext not in allowed_extensions:
        raise FileValidationError(
            f"Unsupported file type: {file_ext}. "
            f"Allowed types: {', '.join(allowed_extensions)}"
        )


def validate_file_size(file: BinaryIO, max_size: int) -> int:
    """
    Validate that the file size is within limits.

    Args:
        file: File object to validate
        max_size: Maximum allowed file size in bytes

    Returns:
        File size in bytes

    Raises:
        FileValidationError: If file is too large
    """
    # Get file size
    file.seek(0, 2)  # Seek to end
    file_size = file.tell()
    file.seek(0)  # Reset to beginning

    if file_size > max_size:
        max_mb = max_size / (1024 * 1024)
        actual_mb = file_size / (1024 * 1024)
        raise FileValidationError(
            f"File too large: {actual_mb:.2f} MB. Maximum size: {max_mb:.2f} MB"
        )

    if file_size == 0:
        raise FileValidationError("File is empty")

    return file_size


def parse_excel(file: BinaryIO, filename: str) -> ParsedDocument:
    """
    Parse an Excel file and extract data with column names preserved.

    Args:
        file: File-like object containing Excel data
        filename: Name of the file being parsed

    Returns:
        ParsedDocument containing text chunks with metadata

    Raises:
        FileValidationError: If file validation fails
        Exception: If parsing fails
    """
    logger.info(f"Parsing Excel file: {filename}")

    # Validate file type
    validate_file_type(filename, EXCEL_EXTENSIONS)

    # Validate file size
    file_size = validate_file_size(file, MAX_EXCEL_SIZE)

    try:
        # Read Excel file
        excel_file = pd.ExcelFile(file)
        chunks = []
        chunk_index = 0

        # Process each sheet
        for sheet_name in excel_file.sheet_names:
            logger.debug(f"Processing sheet: {sheet_name}")

            # Read the sheet
            df = pd.read_excel(excel_file, sheet_name=sheet_name)

            # Skip empty sheets
            if df.empty:
                logger.warning(f"Sheet '{sheet_name}' is empty, skipping")
                continue

            # Get column names
            column_names = df.columns.tolist()
            total_rows = len(df)

            # Process each row
            for row_idx, (_, row) in enumerate(df.iterrows()):
                # Create text representation preserving column names
                row_parts = []
                for col_name, value in row.items():
                    # Skip NaN values
                    if pd.notna(value):
                        row_parts.append(f"{col_name}: {value}")

                # Only create chunk if row has data
                if row_parts:
                    content = " | ".join(row_parts)

                    # Create Excel-specific metadata
                    excel_metadata = ExcelMetadata(
                        sheet_name=str(sheet_name),
                        row_number=row_idx,
                        column_names=column_names,
                        total_rows=total_rows,
                    )

                    # Create chunk with metadata
                    chunk = DocumentChunk(
                        content=content,
                        metadata={
                            "sheet_name": sheet_name,
                            "row_number": row_idx,
                            "column_names": column_names,
                            "total_rows": total_rows,
                            "filename": filename,
                            "file_type": "excel",
                        },
                        chunk_index=chunk_index,
                    )
                    chunks.append(chunk)
                    chunk_index += 1

        if not chunks:
            raise FileValidationError("No data found in Excel file")

        logger.info(f"Successfully parsed Excel file: {filename} ({len(chunks)} chunks)")

        # Create parsed document
        return ParsedDocument(
            filename=filename,
            file_type=FileType.EXCEL,
            chunks=chunks,
            total_chunks=len(chunks),
            file_size_bytes=file_size,
        )

    except FileValidationError:
        raise
    except Exception as e:
        logger.error(f"Error parsing Excel file {filename}: {e}")
        raise Exception(f"Failed to parse Excel file: {str(e)}")


def parse_word(file: BinaryIO, filename: str) -> ParsedDocument:
    """
    Parse a Word document and extract text content.

    Args:
        file: File-like object containing Word document data
        filename: Name of the file being parsed

    Returns:
        ParsedDocument containing text chunks with metadata

    Raises:
        FileValidationError: If file validation fails
        Exception: If parsing fails
    """
    logger.info(f"Parsing Word file: {filename}")

    # Validate file type
    validate_file_type(filename, WORD_EXTENSIONS)

    # Validate file size
    file_size = validate_file_size(file, MAX_WORD_SIZE)

    try:
        # Read Word document
        doc = Document(file)
        chunks = []

        # Get total paragraphs for metadata
        total_paragraphs = len(doc.paragraphs)

        # Process each paragraph
        for para_idx, paragraph in enumerate(doc.paragraphs):
            # Skip empty paragraphs
            text = paragraph.text.strip()
            if not text:
                continue

            # Check if paragraph has special formatting
            has_formatting = False
            if paragraph.runs:
                has_formatting = any(
                    run.bold or run.italic or run.underline
                    for run in paragraph.runs
                )

            # Create Word-specific metadata
            word_metadata = WordMetadata(
                paragraph_number=para_idx,
                total_paragraphs=total_paragraphs,
                has_formatting=has_formatting,
            )

            # Create chunk with metadata
            chunk = DocumentChunk(
                content=text,
                metadata={
                    "paragraph_number": para_idx,
                    "total_paragraphs": total_paragraphs,
                    "has_formatting": has_formatting,
                    "filename": filename,
                    "file_type": "word",
                },
                chunk_index=len(chunks),
            )
            chunks.append(chunk)

        if not chunks:
            raise FileValidationError("No text content found in Word document")

        logger.info(f"Successfully parsed Word file: {filename} ({len(chunks)} chunks)")

        # Create parsed document
        return ParsedDocument(
            filename=filename,
            file_type=FileType.WORD,
            chunks=chunks,
            total_chunks=len(chunks),
            file_size_bytes=file_size,
        )

    except FileValidationError:
        raise
    except Exception as e:
        logger.error(f"Error parsing Word file {filename}: {e}")
        raise Exception(f"Failed to parse Word document: {str(e)}")


def parse_file(file: BinaryIO, filename: str) -> ParsedDocument:
    """
    Parse a file based on its extension.

    Args:
        file: File-like object to parse
        filename: Name of the file

    Returns:
        ParsedDocument containing parsed content

    Raises:
        FileValidationError: If file type is not supported or validation fails
    """
    file_ext = Path(filename).suffix.lower()

    if file_ext in EXCEL_EXTENSIONS:
        return parse_excel(file, filename)
    elif file_ext in WORD_EXTENSIONS:
        return parse_word(file, filename)
    else:
        raise FileValidationError(
            f"Unsupported file type: {file_ext}. "
            f"Supported types: {', '.join(EXCEL_EXTENSIONS | WORD_EXTENSIONS)}"
        )